"""
Feature 1: Universal Document Ingestion
Reads PDFs, Word files, scanned images, screenshots, camera photos
Handles messy layouts, OCR for non-digital documents
Automatic language detection
"""
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import re
from typing import Tuple
from pathlib import Path

class DocumentIngestion:
    """Universal document ingestion with OCR support"""
    
    SUPPORTED_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp', '.tiff', '.bmp'}
    SUPPORTED_DOC_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}
    
    def __init__(self):
        self.tesseract_langs = "eng+urd+hin+ara+spa+fra+deu+chi_sim"
    
    async def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from any supported document format
        Returns: (extracted_text, detected_language)
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif ext in {'.doc', '.docx'}:
            return await self._extract_from_word(file_path)
        elif ext in self.SUPPORTED_IMAGE_EXTENSIONS:
            return await self._extract_from_image(file_path)
        elif ext == '.txt':
            return await self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, str]:
        """Extract text from PDF, with OCR fallback for scanned documents"""
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            
            # If page has very little text, it might be scanned
            if len(text.strip()) < 50:
                # Convert page to image for OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                ocr_text = pytesseract.image_to_string(img, lang=self.tesseract_langs)
                text_parts.append(ocr_text)
            else:
                text_parts.append(text)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        language = self._detect_language(full_text)
        
        return full_text, language
    
    async def _extract_from_word(self, file_path: str) -> Tuple[str, str]:
        """Extract text from Word documents"""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            language = self._detect_language(full_text)
            return full_text, language
            
        except ImportError:
            # Fallback: try to read as plain text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text, self._detect_language(text)
    
    async def _extract_from_image(self, file_path: str) -> Tuple[str, str]:
        """Extract text from images using OCR"""
        img = Image.open(file_path)
        
        # Preprocess image for better OCR
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # OCR with multiple language support
        text = pytesseract.image_to_string(img, lang=self.tesseract_langs)
        language = self._detect_language(text)
        
        return text, language
    
    async def _extract_from_text(self, file_path: str) -> Tuple[str, str]:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        language = self._detect_language(text)
        return text, language
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on character analysis"""
        if not text:
            return "en"
        
        # Check for Arabic script
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        if arabic_chars > len(text) * 0.1:
            return "ar"
        
        # Check for Urdu (uses Arabic script + specific chars)
        urdu_specific = sum(1 for c in text if c in 'ٹڈڑںھہیے')
        if urdu_specific > 10:
            return "ur"
        
        # Check for Hindi (Devanagari)
        hindi_chars = sum(1 for c in text if '\u0900' <= c <= '\u097F')
        if hindi_chars > len(text) * 0.1:
            return "hi"
        
        # Check for Chinese
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        if chinese_chars > len(text) * 0.1:
            return "zh"
        
        # Default to English
        return "en"
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        return text.strip()

document_ingestion = DocumentIngestion()
